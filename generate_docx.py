import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_loom_script():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Loom Video Presentation Script & Interview Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ---------------------------------------------------------
    # PART 1: LOOM VIDEO SCRIPT
    # ---------------------------------------------------------
    doc.add_heading('Part 1: Loom Video Presentation Script', level=1)
    doc.add_paragraph("Objective: Explain your project perfectly, demonstrating both technical depth and product understanding.")
    
    doc.add_heading('Intro (0:00 - 0:30)', level=2)
    doc.add_paragraph(
        '"Hi everyone, I\'m [Your Name], and today I want to walk you through a highly specialized product I built: '
        'A Deterministic 2D Jewelry Compositing Pipeline embedded within a modern MERN-stack and Python application. '
        'The core problem this solves is that jewelry brands struggle to create realistic lifestyle imagery of their products on models '
        'without expensive photoshoots. My platform allows users to dynamically composite flat product images onto high-fidelity, '
        'AI-enhanced model templates with pixel-perfect realism in real time."'
    )
    
    doc.add_heading('Demo & Flow (0:30 - 1:30)', level=2)
    doc.add_paragraph(
        '"Let\'s see it in action. Here on the frontend—built with Next.js and Tailwind CSS—the user selects a variation type, '
        'like \'Model Wearing (Front View)\'. When I hit generate, the backend orchestration takes over. '
        'Instead of just pasting an image, my Python Flask backend performs several intelligent computer vision steps. '
        'First, it uses background removal (rembg/U-2-Net) to isolate the jewelry. Next, it dynamically calculates geometric placement '
        'using precise coordinate metadata so the necklace rests exactly on the model\'s collarbone. '
        'Finally, it builds a dynamic Gaussian occlusion mask to naturally fade the chains as they wrap around the neck, '
        'preventing harsh cutoffs and maintaining anatomical realism."'
    )
    
    doc.add_heading('Architecture & Technical Highlights (1:30 - 2:30)', level=2)
    doc.add_paragraph(
        '"Architecturally, this is a decoupled system. The frontend handles auth and state management seamlessly. '
        'But the real magic is the Python microservice. I intentionally chose a deterministic OpenCV and PIL approach over '
        'pure Generative AI for the final compositing step. Why? Because Generative AI is unpredictable and often hallucinates or alters '
        'the actual product design, which is unacceptable for luxury jewelry brands. By using deterministic geometric scaling, alpha matting, '
        'and dynamic occlusion masks, I guarantee the product remains 100% accurate while still looking photorealistic on the AI templates."'
    )
    
    doc.add_heading('Outro (2:30 - 3:00)', level=2)
    doc.add_paragraph(
        '"In conclusion, building this required balancing complex computer vision mathematics with a sleek, responsive user interface. '
        'It solves a real business problem with a highly optimized architecture. Thank you for watching, and I\'d love to dive deeper into '
        'the code with you."'
    )

    # ---------------------------------------------------------
    # PART 2: ARCHITECTURE & FLOW
    # ---------------------------------------------------------
    doc.add_heading('Part 2: Architecture & Flow', level=1)
    
    doc.add_heading('High-Level Architecture', level=2)
    doc.add_paragraph(
        '1. Frontend (Next.js / React / Tailwind): Handles user auth, dashboard UI, and async polling for task status.\n'
        '2. Database / Auth (Supabase / PostgreSQL): Manages users, JWT authentication, and task history.\n'
        '3. Backend (Python / Flask): Acts as the orchestration layer for image processing.\n'
        '4. Vision Pipeline (OpenCV / PIL / RemBG): Executes the core business logic (extraction, scaling, warping, masking, compositing).'
    )
    
    doc.add_heading('The Compositing Pipeline Flow', level=2)
    doc.add_paragraph(
        'Step 1: Background Removal - Extracts the jewelry using U-2-Net (rembg) and applies morphological operations (dilation/erosion) to preserve delicate chains.\n'
        'Step 2: Template Resolution - Loads the requested high-fidelity template and retrieves its specific metadata (anchor_x, anchor_y, jawline_y, target_width).\n'
        'Step 3: Geometric Scaling - Scales the extracted jewelry precisely to the target_width to match the template\'s zoom ratio.\n'
        'Step 4: Occlusion Masking - Draws a dynamic neck column (e.g., ±300px wide) and chest ellipse using OpenCV, then applies a heavy Gaussian blur to organically fade the chain ends.\n'
        'Step 5: Shadow & Blending - Generates contact drop shadows and blends the alpha layers.'
    )

    # ---------------------------------------------------------
    # PART 3: ENGINEERING DECISIONS
    # ---------------------------------------------------------
    doc.add_heading('Part 3: Critical Engineering Decisions', level=1)
    
    doc.add_paragraph('Decision 1: Deterministic Math vs. Pure Generative AI')
    p = doc.add_paragraph('Why: ')
    p.add_run('Generative AI (like Stable Diffusion inpaint) is amazing, but it hallucinates. If a jewelry brand uploads a $5,000 necklace, the AI might add an extra diamond or alter the chain pattern. By using OpenCV/PIL for the actual product placement (Deterministic), we guarantee 100% product fidelity while using AI solely for generating the background/model templates.')
    
    doc.add_paragraph('Decision 2: Dynamic Occlusion Masking vs. Static Eraser')
    p = doc.add_paragraph('Why: ')
    p.add_run('Instead of permanently erasing the top of the jewelry image, I dynamically generate a Gaussian blur mask based on the jawline coordinates. This means the same necklace can be placed on a zoomed-out template or a macro-closeup template, and the chains will always fade organically around the neck without hard, pixelated cutoffs.')
    
    doc.add_paragraph('Decision 3: Decoupled Architecture')
    p = doc.add_paragraph('Why: ')
    p.add_run('Image processing blocks the main thread in Python. By keeping the frontend separate and using an async task-polling pattern (or background threads), the UI remains responsive while the heavy OpenCV matrix calculations run on the server.')

    # ---------------------------------------------------------
    # PART 4: INTERVIEW Q&A
    # ---------------------------------------------------------
    doc.add_heading('Part 4: Interview Q&A (To Improve Selection Chances)', level=1)
    
    doc.add_heading('Q1: What was the hardest technical challenge in this project?', level=2)
    doc.add_paragraph(
        'A: "The biggest challenge was handling the neck occlusion on the closeup templates. When I scaled the necklace up to fit a zoomed-in macro shot, the chains were getting artificially cut off by my neck-protection mask. I had to refactor the mask generation to dynamically widen the protection column (from 180px to 300px) based on the template\'s scale, ensuring the chains perfectly traced the zoomed-in neck without fading too early."'
    )
    
    doc.add_heading('Q2: How did you optimize the image processing speed?', level=2)
    doc.add_paragraph(
        'A: "I optimized it by aggressively cropping the bounding box immediately after background extraction. By stripping away all empty transparent space, all subsequent OpenCV matrix operations (scaling, blurring, warping) were performed on a much smaller NumPy array, drastically reducing CPU cycles and memory footprint."'
    )
    
    doc.add_heading('Q3: Why use OpenCV instead of just CSS or Canvas on the frontend?', level=2)
    doc.add_paragraph(
        'A: "While CSS can overlap images, we needed true pixel-level manipulation: creating dynamic alpha masks, running connected component analysis to clean up noise, and applying perspective warps. These require heavy matrix operations that are far more efficient and robust in a dedicated Python/OpenCV backend environment."'
    )
    
    doc.add_heading('Q4: How would you scale this if you had 10,000 users?', level=2)
    doc.add_paragraph(
        'A: "Currently, it runs synchronously or in simple threads. To scale, I would implement a message broker like Redis/Celery. The frontend would drop a task into the queue, a pool of worker nodes would process the OpenCV compositing concurrently, and the result would be uploaded to an S3 bucket with a presigned URL returned to the client."'
    )

    # Save
    doc.save('Loom_Presentation_Script.docx')

if __name__ == '__main__':
    create_loom_script()
    print("Docx created successfully.")
